"""
Generate Awaaz Admin Web Project Handover Document (DOCX + PDF)
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from datetime import date
import os

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCX_PATH = os.path.join(PROJECT_ROOT, "Awaaz_Admin_Web_Project_Handover_Document.docx")
PDF_PATH = os.path.join(PROJECT_ROOT, "Awaaz_Admin_Web_Project_Handover_Document.pdf")
TODAY = date.today().strftime("%B %d, %Y")

SECTIONS = [
    ("1. Project Overview", [
        ("Project Name", "Awaaz Operator — Admin Web Panel (package: eagle_eye_admin)"),
        ("Purpose and Goals", (
            "The Awaaz Operator Admin Web Panel is a Flutter-based web application that enables "
            "authorized operators and administrators to manage content and users for the Awaaz mobile "
            "platform. Operators create and moderate Events, Rescue posts, and General posts; review "
            "user reports; manage app users; and (for Super Admins) approve or reject admin registrations. "
            "The panel is optimized for desktop browsers with responsive support for tablet and mobile widths."
        )),
        ("Target Platforms", (
            "• Primary: Web (Flutter Web) — deployed via Firebase Hosting\n"
            "• Secondary (codebase supports): Android, iOS, macOS, Windows\n"
            "• Admin handover scope: Web deployment and browser-based operator workflows"
        )),
        ("Key Features Summary", (
            "• Email/password and Google Sign-In authentication\n"
            "• Multi-module dashboard: Events, Rescue, Reports, User Management, Super Admin, General Posts\n"
            "• Event/Rescue/General post creation, drafts, timelines, and media attachments\n"
            "• Location-based filtering with radius and calendar filters\n"
            "• Category and subcategory management\n"
            "• Reported posts, comments, and users moderation\n"
            "• App user search, profile view, block/unblock\n"
            "• Super Admin approval workflow for admin users\n"
            "• Push notifications (Firebase Cloud Messaging) on supported platforms\n"
            "• Dark-themed operator UI (1920×1080 design baseline)"
        )),
    ]),
    ("2. Technology Stack", [
        ("Flutter Version", "SDK constraint: ^3.5.4 (pubspec.yaml). Verified environment: Flutter 3.44.1 (stable)"),
        ("Dart Version", "SDK constraint: ^3.5.4. Verified environment: Dart 3.12.1"),
        ("State Management", (
            "GetX (get: 4.6.6) — primary pattern for Admin Web:\n"
            "• GetxController classes in lib/controller/\n"
            "• GetBuilder for UI rebuilds\n"
            "• Rx types (e.g., RxBool) for reactive state in select controllers\n"
            "• Get.put / Get.find for dependency injection\n"
            "Note: Legacy consumer-app code under lib/presentation/ uses Bloc/Cubit but is NOT part of the active Admin Web flow."
        )),
        ("Key Packages and Purpose", (
            "• go_router (^7.0.0) — Declarative routing with GetMaterialApp.router\n"
            "• dio (^5.4.3+1) — HTTP client (secondary; primary API via logic_go_network)\n"
            "• logic_go_network (Git) — RestClient wrapper, APIType (public/protected), Failure handling\n"
            "• get_storage (^2.1.1) — Local persistence (JWT token, user prefs, tab state)\n"
            "• flutter_screenutil (^5.9.3) — Responsive scaling (design size 1920×1080)\n"
            "• firebase_core, firebase_auth, firebase_messaging — Firebase integration\n"
            "• google_sign_in — Google OAuth for admin login\n"
            "• flutter_svg, lottie — UI assets and animations\n"
            "• cached_network_image — Network image caching\n"
            "• file_picker, image_picker, image_picker_web — Media upload (web-aware)\n"
            "• geolocator, geocoding — Location services\n"
            "• table_calendar — Date filtering\n"
            "• dropdown_button2 — Enhanced dropdowns\n"
            "• toastification — Toast notifications\n"
            "• connectivity_plus, internet_connection_checker — Network connectivity\n"
            "• intl — Date/number formatting\n"
            "• video_player, get_thumbnail_video — Video preview and thumbnails\n"
            "• flutter_gen_runner — Auto-generated asset references (lib/gen/)"
        )),
    ]),
    ("3. Project Architecture", [
        ("Architecture Pattern", (
            "Feature-oriented MVC with Repository pattern (not full Clean Architecture).\n\n"
            "Layers:\n"
            "• Presentation: lib/page/ (screens + components), lib/widget/ (shared UI)\n"
            "• Controller: lib/controller/ (GetX controllers — business logic & UI state)\n"
            "• Data: lib/api/repository/ (API calls), lib/model/ (DTOs)\n"
            "• Services: lib/service/ (storage, location)\n"
            "• Routing: lib/route/ (GoRouter config, navigation helpers)\n"
            "• Theme/Utils: lib/theme/, lib/utils/"
        )),
        ("Folder Structure Strategy", (
            "Admin Web code is organized by feature under lib/page/ with matching controllers "
            "and repositories. Shared widgets live in lib/widget/. API endpoints are centralized "
            "in lib/api/api_endpoint.dart. Cross-cutting concerns (auth headers, base URL) are in "
            "lib/utils/app_constants.dart and lib/api/api_endpoint.dart."
        )),
        ("Code Organization Notes", (
            "The repository contains legacy consumer-app modules under lib/presentation/, lib/core/, "
            "and lib/data/ (Bloc/Cubit architecture). These directories are remnants from a shared "
            "monorepo and are NOT wired into main.dart or AppRoutes for the Admin Web application. "
            "Future maintainers should treat lib/page/, lib/controller/, lib/api/, and lib/model/ "
            "as the authoritative Admin Web codebase."
        )),
    ]),
    ("4. Folder Structure Breakdown", [
        ("lib/", "Root application source. Entry point: main.dart"),
        ("lib/page/", (
            "Feature screens and UI components:\n"
            "• dashboard_screen/ — Main shell with side navigation\n"
            "• event_screen/ — Event CRUD, drafts, details, attachments\n"
            "• rescue_screen/ — Rescue posts and updates\n"
            "• general_screen/ — General post management\n"
            "• reports_screen/ — Moderation of reported content\n"
            "• user_management_screen/ — App user listing and actions\n"
            "• super_admin_dashboard/ — Admin approval workflow\n"
            "• login_screen/, register_screen/, forgot_password_screen.dart/ — Auth flows\n"
            "• splash_screen/, onboarding_screen/ — App entry and first-time setup"
        )),
        ("lib/controller/", "17 GetX controllers — one per major feature (event, rescue, login, dashboard, etc.)"),
        ("lib/api/", (
            "• api_endpoint.dart — All REST endpoint path constants and auth headers\n"
            "• repository/ — Auth, Event, Rescue, General, Report, User, SuperAdmin repositories\n"
            "• interceptor/token_interceptor.dart — Dio Bearer token interceptor (currently not attached in ApiControllerMain)\n"
            "• connectivity_check/ — Network connectivity monitoring"
        )),
        ("lib/model/", "22 JSON-serializable data models (events, users, reports, categories, etc.)"),
        ("lib/service/", (
            "• storage/ — GetStorage wrapper (StorageService, StorageKeys)\n"
            "• location/ — Location service helpers"
        )),
        ("lib/widget/", "Shared UI: drawer, buttons, dialogs, text fields, toast, network image loader"),
        ("lib/theme/", "AppColors, ProgressLoader, theme constants"),
        ("lib/route/", "app_route.dart (GoRouter), navigator_route.dart (navigation helpers)"),
        ("lib/utils/", "app_constants.dart (base URLs), responsive_utils.dart, app_functions.dart, location_utils.dart"),
        ("lib/gen/", "Auto-generated assets.gen.dart and fonts.gen.dart (flutter_gen)"),
        ("lib/core/", "Not Available for Admin Web — legacy consumer app shared widgets/utilities"),
        ("lib/presentation/", "Not Available for Admin Web — legacy consumer app Bloc/Cubit screens"),
        ("lib/data/", "Not Available for Admin Web — legacy consumer app models"),
        ("lib/domain/", "Not Available — Clean Architecture domain layer not implemented"),
    ]),
    ("5. Setup & Installation Guide", [
        ("Prerequisites", (
            "• Flutter SDK 3.5.4 or higher (3.44.1 recommended)\n"
            "• Dart SDK 3.5.4+\n"
            "• IDE: Android Studio, VS Code, or Cursor with Flutter/Dart extensions\n"
            "• Git (for logic_go_network dependency from Bitbucket)\n"
            "• Chrome or Edge browser for web development\n"
            "• Firebase CLI (for web deployment): npm install -g firebase-tools\n"
            "• Node.js (for Firebase CLI)"
        )),
        ("Install Dependencies", (
            "1. Clone the repository\n"
            "2. cd to project root (Awaaz/)\n"
            "3. Run: flutter pub get\n"
            "4. Run code generation (if assets changed): dart run build_runner build --delete-conflicting-outputs\n"
            "5. Ensure Bitbucket access for logic_go_network Git dependency"
        )),
        ("Run the Project (Web)", (
            "Development:\n"
            "  flutter run -d chrome\n\n"
            "Release build (local):\n"
            "  flutter build web --web-renderer html --release\n\n"
            "Serve locally:\n"
            "  cd build/web && python -m http.server 8080"
        )),
        ("Environment Configuration", (
            "• Base API URL: lib/utils/app_constants.dart\n"
            "    - Admin API: https://awaazeye.com/admin/v1\n"
            "    - User API: https://awaazeye.com/api/v1\n"
            "• Firebase: lib/firebase_options.dart (auto-generated via FlutterFire CLI)\n"
            "• Web Firebase config: web/index.html (Firebase JS SDK)\n"
            "• Google Sign-In: web/index.html meta tag (google-signin-client_id)\n"
            "• config/secrets.json — present in repo (verify contents are not committed to public repos)\n"
            "• android/secrets.properties, ios/Flutter/Secrets.xcconfig — platform secrets\n"
            "• Radar.io API key: hardcoded in dashboard_controller.dart (location autocomplete) — should be moved to env config\n"
            "• No .env file pattern currently implemented; URLs and keys are in source files"
        )),
    ]),
    ("6. Build & Deployment", [
        ("Web Build", (
            "Command (from README):\n"
            "  flutter build web --web-renderer html --release\n\n"
            "Output directory: build/web/\n"
            "Renderer: HTML (set in web/index.html: window.flutterWebRenderer = \"html\")"
        )),
        ("Firebase Hosting Deployment", (
            "1. flutter build web --web-renderer html --release\n"
            "2. firebase login\n"
            "3. firebase deploy --only hosting\n\n"
            "Configuration:\n"
            "• firebase.json — public: build/web, SPA rewrites to index.html\n"
            "• .firebaserc — project: eagle-eye-oprator"
        )),
        ("APK / IPA Builds", (
            "Android APK:\n"
            "  flutter build apk --release\n\n"
            "Android App Bundle:\n"
            "  flutter build appbundle --release\n\n"
            "iOS IPA:\n"
            "  flutter build ipa --release\n"
            "(Requires macOS, Xcode, Apple Developer account)\n\n"
            "Note: Admin Web primary delivery is browser-based; mobile builds exist but are secondary."
        )),
        ("Release Configuration", (
            "• App version: pubspec.yaml — version: 1.0.0+1\n"
            "• Android: android/app/build.gradle (applicationId, signing)\n"
            "• iOS: ios/Runner/Info.plist, Xcode project settings\n"
            "• Web title: \"Awaaz Operator\" (main.dart, web/index.html)\n"
            "• Firebase project ID: eagle-eye-oprator"
        )),
    ]),
    ("7. State Management Explanation", [
        ("How State Is Handled", (
            "Admin Web uses GetX controllers as the single source of truth per feature. "
            "Controllers extend GetxController and call update() to trigger GetBuilder rebuilds. "
            "Some reactive fields use RxBool/Rx types with .obs. Controllers are registered "
            "via Get.put() at screen level or dashboard initialization."
        )),
        ("Controller Lifecycle", (
            "1. Screen initState or GetBuilder initState calls Get.put(Controller)\n"
            "2. Controller fetches data via Repository\n"
            "3. Controller updates local lists/fields and calls update()\n"
            "4. GetBuilder rebuilds bound UI widgets\n"
            "5. User actions call controller methods → API → update() → UI refresh"
        )),
        ("Example Flow: Login", (
            "UI (LoginScreen)\n"
            "  → User taps Login\n"
            "  → LoginController.loginWithEmailNPasssword(context, loader)\n"
            "  → AuthRepository.loginWithEmail() via restClient.post(APIType.public)\n"
            "  → ResponseModel parsed → StorageService.storeToken()\n"
            "  → NavigatorRoute.navigateToRemoveUntil(AppRoutes.event, context)\n"
            "  → DashboardScreen loads with EventController, DashboardController, etc."
        )),
        ("Example Flow: Event List", (
            "UI (EventDashboardView)\n"
            "  → GetBuilder initState: eventController.getEventsList()\n"
            "  → EventRepository.filterEvents() → REST API\n"
            "  → eventController.filterEvents populated → update()\n"
            "  → EventPostCardView renders list with pagination (pageNumber)"
        )),
        ("Bloc/Cubit", "Not used in Admin Web active codebase. Present only in legacy lib/presentation/ modules."),
    ]),
    ("8. API Integration", [
        ("API Structure", (
            "• Primary client: logic_go_network RestClient (initialized in ApiControllerMain.init())\n"
            "• Base URL: https://awaazeye.com/admin/v1\n"
            "• Endpoint paths: lib/api/api_endpoint.dart\n"
            "• Repositories: lib/api/repository/*.dart\n"
            "• Response wrapper: ResponseModel (status, message, body)\n"
            "• Secondary Dio instance in ApiControllerMain (baseUrl: https://awaazeye.com) — used for ad-hoc calls"
        )),
        ("Base URL Handling", (
            "lib/utils/app_constants.dart:\n"
            "  String baseUrl = 'https://awaazeye.com/admin/v1';\n"
            "  String userBaseUrl = 'https://awaazeye.com/api/v1';\n\n"
            "Repositories concatenate: \"$baseUrl/${ApiEndpoint.<endpoint>}\""
        )),
        ("Authentication Flow", (
            "1. POST auth/login/email — returns JWT token and user object\n"
            "2. Token stored via StorageService.storeToken()\n"
            "3. Protected requests: requestHeader(APIType.protected) adds Bearer token\n"
            "4. Google Sign-In: Firebase Auth → POST auth/login-and-register/google\n"
            "5. Token expiry: AppFunctions.onTokenExpire() clears storage, preserves saved credentials, redirects to login\n"
            "6. Splash screen checks StorageService.getToken() → login or dashboard"
        )),
        ("Error Handling Strategy", (
            "• Repository methods catch Failure from logic_go_network\n"
            "• Error strings parsed from exception message\n"
            "• Returns ResponseModel(status: false, message: error)\n"
            "• UI shows toast via showToast() with AppColors.red for errors\n"
            "• 401 handling partially implemented in AuthRepository (onTokenExpire commented in login)\n"
            "• ProgressLoader shown/hidden around async operations\n"
            "• NetworkConnectionService monitors connectivity (limited UI feedback)"
        )),
        ("Key API Modules", (
            "Auth: login, register, forgot-password, OTP, reset-password, Google sign-in\n"
            "Events: CRUD, filter, drafts, timeline, file upload, status updates\n"
            "Rescue: rescue updates, pending count, status updates\n"
            "General: general post CRUD and drafts\n"
            "Reports: reported posts, comments, users; delete comment; block user\n"
            "Users: app user list, profile, FCM token update, radius update\n"
            "Super Admin: admin list, approve/reject admin status\n"
            "Categories/Reactions: list, create category/subcategory"
        )),
    ]),
    ("9. Key Modules / Features", [
        ("Authentication Module", {
            "desc": "Email/password login, registration, forgot password (OTP), reset password, Google Sign-In.",
            "files": "lib/page/login_screen/, register_screen/, forgot_password_screen.dart/, lib/controller/login_controller.dart, register_controller.dart, forgot_password_controller.dart, lib/api/repository/auth_repository.dart",
            "flow": "Splash → Login/Register → Token stored → Dashboard. Forgot password: email → OTP → reset."
        }),
        ("Dashboard Shell", {
            "desc": "Main container with side drawer navigation and tab switching between modules.",
            "files": "lib/page/dashboard_screen/dashboard_screen.dart, lib/widget/drawer.dart, lib/controller/dashboard_controller.dart",
            "flow": "DashboardController.selectedTab (0=Event, 1=Rescue, 2=Reports, 3=User Mgmt, 4=Super Admin, 5=General). Desktop shows permanent drawer; mobile uses hamburger menu."
        }),
        ("Events Module", {
            "desc": "Create, edit, filter, draft, publish events with categories, location, timeline, and media attachments.",
            "files": "lib/page/event_screen/*, lib/controller/event_controller.dart, event_detail_controller.dart, lib/api/repository/event_repository.dart",
            "flow": "Event list (filter by location/radius/calendar) → Create/Edit → Preview → Publish. Sub-tabs: Active posts (1.1), Drafts. Routes: /dashboard, /event/attachEvent, /event/eventDetails."
        }),
        ("Rescue Module", {
            "desc": "Manage rescue posts, updates, drafts, and timeline entries for emergency/rescue content.",
            "files": "lib/page/rescue_screen/*, lib/controller/rescue_controller.dart, rescue_detail_controller.dart, lib/api/repository/rescue_repository.dart",
            "flow": "Rescue dashboard → Create rescue → Attach files/timeline → Moderate rescue updates (pending/approved)."
        }),
        ("General Posts Module", {
            "desc": "General content posts with categories, drafts, and detail views.",
            "files": "lib/page/general_screen/*, lib/controller/general_controller.dart, general_detail_controller.dart, lib/api/repository/general_repository.dart",
            "flow": "General dashboard → Create/filter posts → Draft management → Detail view at /general/generalDetails."
        }),
        ("Reports Module", {
            "desc": "Review and act on reported posts, comments, and users.",
            "files": "lib/page/reports_screen/*, lib/controller/report_controller.dart, lib/api/repository/report_repository.dart",
            "flow": "Sub-tabs for Posts / Comments / Users → View report details → Delete comment or block user."
        }),
        ("User Management Module", {
            "desc": "Search and manage app users; view profiles; block/unblock users.",
            "files": "lib/page/user_management_screen/*, lib/page/user_profile_screen/, lib/controller/user_management_controller.dart, lib/api/repository/user_manage_repository.dart",
            "flow": "Paginated user list (16 per page) → Search → Open profile dialog → Block/unblock via API."
        }),
        ("Super Admin Module", {
            "desc": "Owner-role admins approve or reject pending admin registrations.",
            "files": "lib/page/super_admin_dashboard/, lib/controller/super_admin_controller.dart, lib/api/repository/super_admin_repository.dart",
            "flow": "Login with role=owner → Super Admin tab → List admins by filter → Approve/Reject. Route: /superAdminDashboard (also accessible as dashboard tab 4)."
        }),
        ("Onboarding Module", {
            "desc": "First-time admin setup including location radius configuration.",
            "files": "lib/page/onboarding_screen/, lib/controller/onboarding_controller.dart",
            "flow": "Post-registration or first login → Onboarding screens → Location/radius setup."
        }),
    ]),
    ("10. UI/UX Structure", [
        ("Navigation Approach", (
            "GoRouter (^7.0.0) integrated via GetMaterialApp.router.\n"
            "Routes defined in lib/route/app_route.dart.\n"
            "Navigation helpers in lib/route/navigator_route.dart (context.push, context.go, context.pop).\n"
            "Legacy GetX routing (GetPage) is commented out.\n"
            "Auth middleware (AuthMiddleware) exists for GetX but is NOT currently applied to GoRouter routes — auth is handled manually in SplashScreen."
        )),
        ("Theme Management", (
            "• Dark theme enforced in main.dart ThemeData\n"
            "• AppColors in lib/theme/colors.dart (black scaffold, drawer #121212, accent blues)\n"
            "• Font family: Poppins (primary), Arvo (secondary)\n"
            "• Material 3 enabled (useMaterial3: true)\n"
            "• ToastificationWrapper for global toast overlay"
        )),
        ("Responsive Design", (
            "• flutter_screenutil with designSize: Size(1920, 1080)\n"
            "• Responsive utility (lib/utils/responsive_utils.dart):\n"
            "    - Mobile: width < 850px\n"
            "    - Tablet: 850px – 1099px\n"
            "    - Desktop: width ≥ 1100px\n"
            "• Desktop: permanent side drawer; Mobile: collapsible drawer via Scaffold.drawer\n"
            "• textScaler fixed to linear(1) to prevent system font scaling issues"
        )),
    ]),
    ("11. Assets & Resources", [
        ("Images", (
            "assets/image/ — logos (aawaz_logo.svg, logo.png), map background, profile placeholders, "
            "category illustrations (fire, police, hospital, traffic, etc.)"
        )),
        ("Icons", "assets/icons/ — SVG and PNG icons (dashboard, search, filter, rescue, reports, user management, etc.)"),
        ("Animations", "assets/animation/ — loader.json, sucess_animation.json (Lottie)"),
        ("Fonts", (
            "font/ — Poppins (Light, Regular, Medium, SemiBold, Bold), Arvo Bold\n"
            "Registered in pubspec.yaml and referenced via ThemeData.fontFamily"
        )),
        ("Localization", (
            "flutter_intl enabled in pubspec.yaml.\n"
            "lib/l10n/intl_en.arb — empty (no strings defined).\n"
            "lib/generated/l10n.dart — generated but unused.\n"
            "Status: Localization infrastructure present; all UI strings are hardcoded in English."
        )),
        ("Asset Code Generation", "flutter_gen generates lib/gen/assets.gen.dart — use Assets.image.*, Assets.icons.* for type-safe asset access."),
    ]),
    ("12. Testing & QA", [
        ("Unit Tests", "Not Available — no unit test files in test/ directory."),
        ("Widget Tests", (
            "test/widget_test.dart exists but contains default Flutter counter smoke test "
            "(expects '0' and '+' icon) — incompatible with current MyApp. Test will fail if run."
        )),
        ("Integration Tests", "Not Available"),
        ("Known Test Coverage", "Effectively 0% for Admin Web. No CI test pipeline configured."),
        ("Manual QA Recommendations", (
            "• Login/logout flows (email + Google)\n"
            "• Event create → draft → publish → filter\n"
            "• Rescue update moderation\n"
            "• Report actions (delete comment, block user)\n"
            "• Super Admin approve/reject\n"
            "• Responsive layout at 1920, 1024, and 375px widths\n"
            "• Firebase Hosting SPA routing (deep links refresh correctly)"
        )),
    ]),
    ("13. Performance Considerations", [
        ("Optimization Techniques", (
            "• cached_network_image for network image caching\n"
            "• Pagination on event lists and user lists (pageNumber increment)\n"
            "• Debounced location search (Timer debounce in DashboardController)\n"
            "• HTML web renderer for broader browser compatibility (vs CanvasKit)\n"
            "• flutter_screenutil for consistent scaling without manual MediaQuery math"
        )),
        ("Memory Management", (
            "• ScrollControllers disposed in controller dispose() methods\n"
            "• Debounce timers cancelled on dispose\n"
            "• GetStorage for lightweight persistence vs heavy local DB\n"
            "• Large file uploads use MultipartFile.fromBytes with size considerations"
        )),
        ("Best Practices Followed", (
            "• Repository pattern separates API from UI\n"
            "• Centralized API endpoints and headers\n"
            "• Type-safe asset generation via flutter_gen\n"
            "• Separation of page components into subfolders"
        )),
        ("Areas for Improvement", (
            "• EventController is very large (~1200+ lines) — consider splitting\n"
            "• Multiple Get.put() calls for same controller across widgets — risk of duplicate instances\n"
            "• Token interceptor commented out in ApiControllerMain"
        )),
    ]),
    ("14. Known Issues / Limitations", [
        ("Bugs / Incomplete Features", (
            "• GoRouter auth redirect is commented out — unauthenticated users can manually navigate to /dashboard\n"
            "• AuthMiddleware designed for GetX routing, not integrated with GoRouter\n"
            "• SplashController navigation logic is commented out (handled in SplashScreen widget instead)\n"
            "• widget_test.dart is outdated and will fail\n"
            "• NavigatorRoute.navigateToReplacement still uses Get.offAndToNamed (inconsistent with GoRouter)\n"
            "• Super Admin dashboard route exists but owner users are redirected to /dashboard (tab 4) instead of dedicated route\n"
            "• Hardcoded Radar.io API key in dashboard_controller.dart (security concern)\n"
            "• Commented hardcoded JWT in api_endpoint.dart requestHeader (must never be enabled in production)"
        )),
        ("Pending Features", (
            "• Full localization (intl_en.arb is empty)\n"
            "• Automated test suite\n"
            "• Environment-based config (.env / flavors for dev/staging/prod)\n"
            "• GoRouter auth guard implementation"
        )),
        ("Technical Debt", (
            "• Legacy consumer app code (lib/presentation/, lib/core/, lib/data/) clutters repository\n"
            "• Mixed navigation APIs (GoRouter + residual GetX navigation)\n"
            "• Typo naming conventions (dailoge, feild, calander) throughout codebase\n"
            "• Duplicate model definitions (lib/model/ vs lib/data/models/)\n"
            "• secrets.json and secrets.properties in working tree — verify .gitignore coverage"
        )),
    ]),
    ("15. Future Improvements", [
        ("Suggested Enhancements", (
            "• Implement GoRouter redirect for authentication and role-based access\n"
            "• Extract environment variables (API URL, Radar key, Firebase config) to secure config\n"
            "• Add comprehensive widget and integration tests\n"
            "• Remove or isolate legacy consumer app code into separate package\n"
            "• Implement proper error boundary and global error handler\n"
            "• Add audit logging for admin actions (approve, block, delete)\n"
            "• Enable CanvasKit or Wasm renderer evaluation for web performance\n"
            "• Complete i18n with intl code generation"
        )),
        ("Scalability Improvements", (
            "• Split large controllers into use-case services\n"
            "• Introduce repository interfaces for testability\n"
            "• Add API response caching layer for categories/reactions\n"
            "• Implement proper state management per module with GetX bindings\n"
            "• Consider feature modules with lazy loading for web bundle size\n"
            "• Set up CI/CD pipeline (build web → Firebase deploy on merge)"
        )),
    ]),
    ("16. Credentials & Access", [
        ("Firebase Project", "Project ID: eagle-eye-oprator | Hosting: Firebase Hosting (build/web)"),
        ("API Base URLs", "Production Admin API: https://awaazeye.com/admin/v1 | User API: https://awaazeye.com/api/v1"),
        ("Google Sign-In (Web)", "Client ID: 1086312563652-****.apps.googleusercontent.com (see web/index.html)"),
        ("Firebase Web API Key", "AIzaSy**** (see web/index.html and lib/firebase_options.dart) — restrict by domain in Firebase Console"),
        ("Radar.io API Key", "prj_live_pk_**** (hardcoded in dashboard_controller.dart) — rotate and move to secrets"),
        ("Admin Access Notes", (
            "• Admin accounts require backend registration and approval (Super Admin for owner role)\n"
            "• Role 'owner' enables Super Admin tab and isSuperAdmin flag in local storage\n"
            "• JWT Bearer token required for all protected API endpoints\n"
            "• Do NOT commit production tokens or API keys to version control"
        )),
        ("Secrets Files", "config/secrets.json, android/secrets.properties, ios/Flutter/Secrets.xcconfig — verify access with DevOps; mask in documentation"),
    ]),
    ("17. Handover Notes", [
        ("Important Developer Notes", (
            "• This is a Flutter Web-first admin panel; always test changes in Chrome at 1920×1080\n"
            "• Use flutter build web --web-renderer html --release before Firebase deploy\n"
            "• logic_go_network is a private Bitbucket dependency — ensure team access\n"
            "• Primary active code paths: main.dart → AppRoutes.router → lib/page/*\n"
            "• Ignore lib/presentation/ for Admin Web unless explicitly migrating features\n"
            "• Run dart run build_runner build after adding/modifying assets\n"
            "• Package name in pubspec is eagle_eye_admin; app title is 'Awaaz Operator'"
        )),
        ("Things to Be Careful About", (
            "• Do not enable commented hardcoded JWT in api_endpoint.dart\n"
            "• Auth is NOT enforced at router level — always implement redirect before production hardening\n"
            "• Get.put() in multiple widgets can cause state inconsistency — prefer Bindings or single put at route level\n"
            "• File uploads differ between web (bytes) and mobile (file path) — test both if supporting mobile\n"
            "• Firebase messaging on web requires service worker (firebase-messaging-sw.js)\n"
            "• Clearing storage on logout must preserve 'keep me logged in' credentials per AppFunctions.onTokenExpire pattern\n"
            "• Before handover to client: rotate any exposed API keys found in source code"
        )),
        ("Document Metadata", f"Generated: {TODAY} | Document Version: 1.0 | Scope: Awaaz Admin Web Panel"),
    ]),
]


def add_docx_content(doc):
    # Title
    title = doc.add_heading("Awaaz Admin Web", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Project Handover Document")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    meta = doc.add_paragraph(f"Version 1.0  |  {TODAY}  |  Confidential")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # TOC placeholder
    doc.add_heading("Table of Contents", level=1)
    for section_title, _ in SECTIONS:
        doc.add_paragraph(section_title, style="List Number")
    doc.add_page_break()

    for section_title, items in SECTIONS:
        doc.add_heading(section_title, level=1)
        for item in items:
            if isinstance(item, tuple) and len(item) == 2:
                key, val = item
                p = doc.add_paragraph()
                run = p.add_run(f"{key}: ")
                run.bold = True
                if isinstance(val, str) and "\n" in val:
                    doc.add_paragraph(val)
                else:
                    p.add_run(str(val))
            elif isinstance(item, dict):
                for mod_name, mod_data in item.items():
                    doc.add_heading(mod_name, level=2)
                    doc.add_paragraph(f"Description: {mod_data['desc']}")
                    doc.add_paragraph(f"Main Files: {mod_data['files']}")
                    doc.add_paragraph(f"Flow: {mod_data['flow']}")
        doc.add_paragraph()


def build_pdf():
    doc = SimpleDocTemplate(
        PDF_PATH,
        pagesize=A4,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="DocTitle", fontSize=22, leading=26, alignment=TA_CENTER,
        textColor=colors.HexColor("#1a1a2e"), spaceAfter=6, fontName="Helvetica-Bold"
    ))
    styles.add(ParagraphStyle(
        name="DocSubtitle", fontSize=14, leading=18, alignment=TA_CENTER,
        textColor=colors.HexColor("#444444"), spaceAfter=4
    ))
    styles.add(ParagraphStyle(
        name="SectionHeader", fontSize=14, leading=18, spaceBefore=14, spaceAfter=8,
        textColor=colors.HexColor("#1a5276"), fontName="Helvetica-Bold"
    ))
    styles.add(ParagraphStyle(
        name="SubHeader", fontSize=11, leading=14, spaceBefore=8, spaceAfter=4,
        textColor=colors.HexColor("#2c3e50"), fontName="Helvetica-Bold"
    ))
    styles.add(ParagraphStyle(
        name="BodyText2", fontSize=9.5, leading=13, alignment=TA_JUSTIFY, spaceAfter=4
    ))
    styles.add(ParagraphStyle(
        name="BulletText", fontSize=9.5, leading=13, leftIndent=12, spaceAfter=2
    ))

    story = []
    story.append(Paragraph("Awaaz Admin Web", styles["DocTitle"]))
    story.append(Paragraph("Project Handover Document", styles["DocSubtitle"]))
    story.append(Paragraph(f"Version 1.0 | {TODAY} | Confidential", styles["DocSubtitle"]))
    story.append(Spacer(1, 0.3 * inch))

    # Route map table
    story.append(Paragraph("Quick Reference — Application Routes", styles["SectionHeader"]))
    route_data = [
        ["Route", "Screen"],
        ["/", "Splash"],
        ["/login", "Login"],
        ["/register", "Register"],
        ["/forgotPassword", "Forgot Password"],
        ["/otpScreen", "OTP Verification"],
        ["/resetPassword", "Reset Password"],
        ["/onBoarding", "Onboarding"],
        ["/dashboard", "Main Dashboard (Events default tab)"],
        ["/rescue", "Rescue Dashboard"],
        ["/event/attachEvent", "Attach Event Media"],
        ["/event/eventDetails", "Event Details"],
        ["/rescue/rescueDetails", "Rescue Details"],
        ["/general/generalDetails", "General Post Details"],
        ["/superAdminDashboard", "Super Admin"],
        ["/event/userProfile/:userId", "User Profile"],
    ]
    t = Table(route_data, colWidths=[2.8 * inch, 3.5 * inch])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a5276")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f4f6f7")]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(t)
    story.append(Spacer(1, 0.2 * inch))

    for section_title, items in SECTIONS:
        story.append(Paragraph(section_title, styles["SectionHeader"]))
        for item in items:
            if isinstance(item, tuple) and len(item) == 2:
                key, val = item
                story.append(Paragraph(f"<b>{key}</b>", styles["SubHeader"]))
                for line in str(val).split("\n"):
                    line = line.strip()
                    if line:
                        if line.startswith("•"):
                            story.append(Paragraph(line, styles["BulletText"]))
                        else:
                            story.append(Paragraph(line, styles["BodyText2"]))
            elif isinstance(item, dict):
                for mod_name, mod_data in item.items():
                    story.append(Paragraph(mod_name, styles["SubHeader"]))
                    story.append(Paragraph(f"<b>Description:</b> {mod_data['desc']}", styles["BodyText2"]))
                    story.append(Paragraph(f"<b>Main Files:</b> {mod_data['files']}", styles["BodyText2"]))
                    story.append(Paragraph(f"<b>Flow:</b> {mod_data['flow']}", styles["BodyText2"]))
        story.append(Spacer(1, 0.1 * inch))

    doc.build(story)


def main():
    doc = Document()
    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_docx_content(doc)
    doc.save(DOCX_PATH)
    print(f"DOCX created: {DOCX_PATH}")

    build_pdf()
    print(f"PDF created: {PDF_PATH}")


if __name__ == "__main__":
    main()
