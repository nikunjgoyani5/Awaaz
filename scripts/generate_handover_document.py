"""
Generate Awaaz Flutter Mobile App Project Handover Document (DOCX + PDF).
"""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


PROJECT_ROOT = Path(__file__).resolve().parent.parent
DOCX_PATH = PROJECT_ROOT / "Awaaz_Flutter_Mobile_App_Project_Handover_Document.docx"
PDF_PATH = PROJECT_ROOT / "Awaaz_Flutter_Mobile_App_Project_Handover_Document.pdf"
TODAY = date.today().strftime("%B %d, %Y")


def set_document_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_cover_page(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Awaaz Flutter Mobile App")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x66)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Project Handover Document")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Document Version: 1.0\n").bold = True
    meta.add_run(f"Prepared On: {TODAY}\n")
    meta.add_run("Application: Awaaz – Real Time City Alert\n")
    meta.add_run("Package Name: eagle_eye\n")
    meta.add_run("App Version: 1.0.13+14\n")
    meta.add_run("Classification: Client & Developer Handover")

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def build_document() -> Document:
    doc = Document()
    set_document_styles(doc)

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    add_cover_page(doc)

    # Table of Contents placeholder
    add_heading(doc, "Table of Contents", 1)
    toc_items = [
        "1. Project Overview",
        "2. Technology Stack",
        "3. Project Architecture",
        "4. Folder Structure Breakdown",
        "5. Setup & Installation Guide",
        "6. Build & Deployment",
        "7. State Management Explanation",
        "8. API Integration",
        "9. Key Modules / Features",
        "10. UI/UX Structure",
        "11. Assets & Resources",
        "12. Testing & QA",
        "13. Performance Considerations",
        "14. Known Issues / Limitations",
        "15. Future Improvements",
        "16. Credentials & Access",
        "17. Handover Notes",
    ]
    add_bullets(doc, toc_items)
    doc.add_page_break()

    # 1. Project Overview
    add_heading(doc, "1. Project Overview", 1)
    add_heading(doc, "Project Name", 2)
    add_paragraph(doc, "Awaaz – Real Time City Alert (internal Dart package name: eagle_eye)")
    add_heading(doc, "Purpose and Goals", 2)
    add_paragraph(
        doc,
        "Awaaz is a community-driven, location-aware mobile application that enables citizens to "
        "share real-time incident alerts, live video updates, and general community posts. The app "
        "helps users stay informed about events happening within a configurable radius (default ~10 km), "
        "interact through comments and reactions, chat with other users, and report or block inappropriate content.",
    )
    add_heading(doc, "Target Platforms", 2)
    add_bullets(
        doc,
        [
            "Android (primary) – Application ID: com.awaazeye.cityalerts, minSdk 26, targetSdk 35",
            "iOS (primary) – Display name: Awaaz, CFBundleShortVersionString 1.2",
            "Web – Scaffold present (web/ folder, Firebase web config); not the primary delivery target",
            "Desktop (Windows/macOS) – Firebase options exist; not actively configured for production release",
        ],
    )
    add_heading(doc, "Key Features Summary", 2)
    add_bullets(
        doc,
        [
            "Guest/device-based login and full authentication (email, Google, Apple Sign-In)",
            "Real-time news feed of incident and general posts with pagination",
            "Interactive Mapbox map showing nearby events",
            "Go Live – capture and publish live incident/general video alerts with categories",
            "Send Alert – targeted alerts to nearby users",
            "News details with comments, reactions, in-area related posts, and reporting",
            "User profiles, edit profile, friends, chat (Socket.IO), and blocked users",
            "Push notifications via OneSignal with deep-link navigation",
            "Branch.io deep linking for shared event URLs",
            "In-app support ticket system",
            "Settings: video quality, notification preferences, account deletion",
            "Google Mobile Ads with Firebase Remote Config-driven ad placement",
            "Firebase Analytics, Crashlytics, Performance Monitoring, and Remote Config",
        ],
    )

    # 2. Technology Stack
    add_heading(doc, "2. Technology Stack", 1)
    add_bullets(
        doc,
        [
            "Flutter Version: 3.44.1 (stable channel)",
            "Dart Version: 3.12.1 (SDK constraint in pubspec.yaml: ^3.5.3)",
            "State Management: flutter_bloc (BLoC/Cubit pattern) with Freezed for immutable states",
            "Secondary utilities: Get package (^4.7.2) imported for limited use; primary navigation uses Navigator 1.0",
            "Networking: logic_go_network (private Bitbucket git dependency) + Dio",
            "Real-time: socket_io_client for chat and live updates",
            "Maps: mapbox_maps_flutter",
            "Local Storage: shared_preferences via PrefService wrapper",
            "Code Generation: build_runner, freezed, flutter_gen_runner",
        ],
    )
    add_heading(doc, "Key Packages and Purpose", 2)
    packages = [
        ("flutter_bloc / bloc / freezed", "State management with immutable Cubit states"),
        ("dio / logic_go_network", "REST API communication with typed APIType (public/protected)"),
        ("socket_io_client", "Real-time chat, typing indicators, unread counts"),
        ("mapbox_maps_flutter", "Interactive map with event annotations"),
        ("firebase_core, firebase_auth, firebase_analytics, firebase_crashlytics, firebase_performance, firebase_remote_config", "Firebase platform services"),
        ("onesignal_flutter", "Push notification delivery and click handling"),
        ("google_sign_in / sign_in_with_apple", "Social authentication"),
        ("flutter_branch_sdk / app_links", "Deep linking and share attribution"),
        ("google_mobile_ads", "Monetization (banner, native, interstitial/open app ads)"),
        ("image_picker, file_picker, camera, chewie, video_player, light_compressor_v2", "Media capture, playback, and compression"),
        ("geolocator, location, geocoding, permission_handler", "Location services and permissions"),
        ("flutter_screenutil", "Responsive UI scaling (design size 430×990)"),
        ("connectivity_plus / internet_connection_checker", "Network connectivity monitoring"),
        ("cached_network_image", "Optimized remote image loading"),
        ("showcaseview", "In-app tutorial overlays"),
        ("lottie", "Loading and success animations"),
        ("youtube_player_flutter", "Embedded tutorial/help videos"),
    ]
    for name, purpose in packages:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{name}: ").bold = True
        p.add_run(purpose)

    # 3. Project Architecture
    add_heading(doc, "3. Project Architecture", 1)
    add_heading(doc, "Architecture Pattern", 2)
    add_paragraph(
        doc,
        "The project follows a layered, feature-oriented architecture inspired by Clean Architecture "
        "principles but without a dedicated domain layer. The structure separates concerns into "
        "presentation (UI + Cubits), data (repositories + models + API), core (shared utilities, "
        "theme, widgets), and services (platform integrations). This is commonly referred to as a "
        "Presentation–Data–Core pattern with BLoC for state management.",
    )
    add_heading(doc, "Code Organization Strategy", 2)
    add_bullets(
        doc,
        [
            "Feature modules live under lib/presentation/ grouped by flow (onboard, authentication, main)",
            "Each feature typically contains a screen widget and a bloc/ subfolder with Cubit + State (Freezed)",
            "API calls are centralized in lib/data/repositories/ (auth_repo.dart, main_repo.dart)",
            "Shared UI components reside in lib/core/widget/",
            "Global Cubits are registered once in lib/core/bloc_provider/bloc_providers.dart",
            "Routes are defined in lib/routes/app_routes.dart with onGenerateRoute",
            "Platform services (notifications, ads, sockets, camera) are isolated in lib/services/",
        ],
    )

    # 4. Folder Structure
    add_heading(doc, "4. Folder Structure Breakdown", 1)
    folders = [
        (
            "lib/",
            "Root application source. Entry point: main.dart. Contains all Dart application code.",
        ),
        (
            "lib/core/",
            "Shared foundation: constants, theme, reusable widgets, utilities (PrefService, MapUtils, AppFunctions), "
            "and global BlocProvider registration.",
        ),
        (
            "lib/core/constant/",
            "App-wide constants: base URLs, event categories, report options, extensions.",
        ),
        (
            "lib/core/theme/",
            "ThemeData (darkTheme), color palette, and text styles.",
        ),
        (
            "lib/core/widget/",
            "Reusable UI components: buttons, text fields, loaders, dialogs, bottom sheets, video players.",
        ),
        (
            "lib/core/utils/",
            "Helper utilities: preferences, map helpers, skeleton loaders, hashtag formatter.",
        ),
        (
            "lib/core/bloc_provider/",
            "MultiBlocProvider configuration for all global Cubits.",
        ),
        (
            "lib/presentation/",
            "All UI screens organized by feature area: onboard/, authentication/, main/.",
        ),
        (
            "lib/data/",
            "Data layer: API controller, endpoints, models, and repositories.",
        ),
        (
            "lib/data/models/",
            "JSON-serializable data models (login, events, chat, profile, notifications, etc.).",
        ),
        (
            "lib/data/repositories/",
            "auth_repo.dart – authentication APIs; main_repo.dart – all other protected/public APIs.",
        ),
        (
            "lib/services/",
            "Platform and third-party integrations: OneSignal, Firebase, ads, sockets, camera, deep linking, connectivity.",
        ),
        (
            "lib/routes/",
            "app_routes.dart (route names + transitions), app_navigation.dart (Navigator helpers).",
        ),
        (
            "lib/gen/",
            "Auto-generated asset and font references (flutter_gen).",
        ),
        (
            "assets/",
            "Static assets: icons/, images/, videos/, fonts/, animation/ (Lottie JSON).",
        ),
        (
            "android/ / ios/",
            "Native platform configuration, permissions, signing, Firebase, Branch, and Mapbox setup.",
        ),
        (
            "test/",
            "Flutter test directory (minimal coverage – see Section 12).",
        ),
        ("lib/features/", "Not Available – features are organized under lib/presentation/ instead."),
        ("lib/domain/", "Not Available – no separate domain/use-case layer."),
    ]
    for name, desc in folders:
        p = doc.add_paragraph()
        p.add_run(f"{name} – ").bold = True
        p.add_run(desc)

    # 5. Setup
    add_heading(doc, "5. Setup & Installation Guide", 1)
    add_heading(doc, "Prerequisites", 2)
    add_bullets(
        doc,
        [
            "Flutter SDK 3.44.x or compatible (Dart ^3.5.3)",
            "Android Studio / VS Code with Flutter & Dart plugins",
            "Xcode (macOS only, for iOS builds)",
            "Git access to Bitbucket for logic_go_network private dependency",
            "Firebase project access (aawaj-52af1)",
            "Mapbox access token",
            "OneSignal, Branch.io, and Google AdMob account credentials",
        ],
    )
    add_heading(doc, "Installation Steps", 2)
    add_numbered(
        doc,
        [
            "Clone the repository to your local machine.",
            "Ensure Flutter is on PATH and run: flutter doctor",
            "From project root, install dependencies: flutter pub get",
            "Run code generation (Freezed / flutter_gen): dart run build_runner build --delete-conflicting-outputs",
            "Place google-services.json (Android) and GoogleService-Info.plist (iOS) if not present.",
            "Configure Mapbox token (see Environment Configuration below).",
            "Connect a device or emulator and run: flutter run",
        ],
    )
    add_heading(doc, "Environment Configuration", 2)
    add_bullets(
        doc,
        [
            "API Environment: Controlled by isLiveMode flag in lib/core/constant/app_constant.dart. "
            "true → https://awaazeye.com; false → http://139.59.32.181:8002",
            "Mapbox Access Token: Passed at build/run time via --dart-define=ACCESS_TOKEN=<your_token>. "
            "Read in main.dart via String.fromEnvironment('ACCESS_TOKEN'). iOS uses $(MAPBOX_ACCESS_TOKEN) in Info.plist.",
            "Firebase: Configured via lib/firebase_options.dart (generated by FlutterFire CLI).",
            "Branch.io Test Mode: Set io.branch.sdk.TestMode to false in AndroidManifest.xml for production (see README.md).",
            "No .env file is used; secrets and toggles are compile-time defines or in-code constants.",
        ],
    )

    # 6. Build & Deployment
    add_heading(doc, "6. Build & Deployment", 1)
    add_heading(doc, "Android", 2)
    add_bullets(
        doc,
        [
            "Debug APK: flutter build apk --debug --dart-define=ACCESS_TOKEN=<token>",
            "Release APK: flutter build apk --release --dart-define=ACCESS_TOKEN=<token>",
            "App Bundle (Play Store): flutter build appbundle --release --dart-define=ACCESS_TOKEN=<token>",
            "Release signing: android/app/build.gradle currently uses signingConfigs.debug for release – replace with production keystore before store submission.",
            "ProGuard/R8: minifyEnabled and shrinkResources enabled for release builds.",
            "ABI filters: armeabi-v7a, arm64-v8a only.",
        ],
    )
    add_heading(doc, "iOS", 2)
    add_bullets(
        doc,
        [
            "Debug: flutter build ios --debug --dart-define=ACCESS_TOKEN=<token>",
            "Release: flutter build ios --release --dart-define=ACCESS_TOKEN=<token>",
            "Archive and upload via Xcode after configuring signing certificates and provisioning profiles.",
            "Set MAPBOX_ACCESS_TOKEN in Xcode build settings or xcconfig.",
        ],
    )
    add_heading(doc, "Web", 2)
    add_paragraph(doc, "Web scaffold exists. Build command: flutter build web. Not the primary production target.")
    add_heading(doc, "Code Generation (CI/CD note)", 2)
    add_paragraph(doc, "Include dart run build_runner build --delete-conflicting-outputs in CI before flutter build.")

    # 7. State Management
    add_heading(doc, "7. State Management Explanation", 1)
    add_paragraph(
        doc,
        "The application uses the BLoC pattern via Cubit classes (flutter_bloc). States are immutable "
        "and defined with Freezed (@freezed) for copyWith, equality, and code generation. All primary "
        "Cubits are provided at the app root through MultiBlocProvider in main.dart.",
    )
    add_heading(doc, "Registered Global Cubits", 2)
    add_bullets(
        doc,
        [
            "OnboardCubit, SplashCubit – onboarding and guest auth",
            "AuthScreenCubit, VerifyOtpCubit, ResetPasswordCubit, ChangePasswordCubit – authentication flows",
            "HomeScreenBlocCubit – bottom navigation and home shell state",
            "NewsScreenBlocCubit, NewsDetailsScreenBlocCubit – feed and detail views",
            "MapScreenCubit, GoLiveCubit, GeneralPostCubit, SendAlertCubit – map and content creation",
            "MyProfileCubit, EditProfileCubit, UserProfileCubit – profile management",
            "AlertsScreenBlocCubit, SearchScreenBlocCubit, SearchLocationCubit",
            "GetSupportCubit, NotificationSettingsCubit, DeleteAccountCubit, BlockedUsersCubit",
        ],
    )
    add_heading(doc, "Example Flow: News Feed (UI → Cubit → API → UI)", 2)
    add_numbered(
        doc,
        [
            "User opens HomeScreen → NewsScreen tab is first page in bottom navigation.",
            "NewsScreen dispatches getNewsEvents('latest') on NewsScreenBlocCubit.",
            "Cubit emits isLoading: true, then calls MainRepository.eventNewsList().",
            "MainRepository uses logic_go_network RestClient with APIType.protected and Bearer token from PrefService.",
            "API endpoint: GET api/v1/event-post/event-post-news?filterType=<type>&page=<n>&limit=<n>",
            "Response parsed into EventNewsModel; Cubit emits updated eventPostList and pagination state.",
            "NewsScreen rebuilds via BlocBuilder/BlocListener showing list or shimmer loader.",
            "On 401 errors, AppFunctions.onTokenExpire() handles session expiry.",
        ],
    )

    # 8. API Integration
    add_heading(doc, "8. API Integration", 1)
    add_heading(doc, "API Structure", 2)
    add_bullets(
        doc,
        [
            "Primary REST client: logic_go_network RestClient initialized in ApiController.init()",
            "Secondary Dio instance in ApiController for legacy/direct calls (base URL hardcoded separately)",
            "Repositories: AuthRepository (auth), MainRepository (all other endpoints)",
            "Response wrapper: ResponseModel { status, message, body }",
            "API versioning prefix: api/v1/",
        ],
    )
    add_heading(doc, "Base URL Handling", 2)
    add_bullets(
        doc,
        [
            "Production: https://awaazeye.com",
            "Development: http://139.59.32.181:8002",
            "Socket URL: same host with trailing slash (baseSocketUrl in app_constant.dart)",
            "Toggle: isLiveMode boolean in app_constant.dart",
        ],
    )
    add_heading(doc, "Authentication Flow", 2)
    add_numbered(
        doc,
        [
            "Splash screen may perform guest login via device ID (api/v1/auth/guest-login).",
            "Full auth supports email/password, mobile OTP, Google, and Apple sign-in.",
            "On success, JWT access token stored in SharedPreferences (PrefService.accessToken).",
            "Protected requests include Authorization: Bearer <token> via requestHeader(APIType.protected).",
            "Token expiry (401) triggers AppFunctions.onTokenExpire() redirecting user to re-authenticate.",
        ],
    )
    add_heading(doc, "Key API Endpoint Groups", 2)
    add_bullets(
        doc,
        [
            "Auth: login, register, OTP verify/resend, forgot/reset/change password, Google/Apple login",
            "User: profile CRUD, location/radius update, push token, block/unblock, delete account",
            "Events: post, draft, reactions, comments, search, map area events, view counts",
            "General posts: add post, draft",
            "Support: list, add support requests",
            "Notifications: user notification list, on/off toggles",
        ],
    )
    add_heading(doc, "Error Handling Strategy", 2)
    add_bullets(
        doc,
        [
            "Repository methods catch logic_go_network Failure exceptions and parse error messages",
            "Errors returned as ResponseModel with status: false and user-facing message",
            "UI Cubits display errors via AppFunctions.showToast() or CommonSnackBar",
            "Connectivity checked via networkConnectionServices before critical operations",
            "Firebase Crashlytics captures unhandled Flutter errors when isLiveMode is true",
        ],
    )

    # 9. Key Modules
    add_heading(doc, "9. Key Modules / Features", 1)
    modules = [
        (
            "Splash & Onboarding",
            "App entry, guest device authentication, name/username/birthdate/profile setup.",
            "lib/presentation/onboard/splash_screen/, lib/presentation/onboard/onboard_screen/, lib/presentation/onboard/new_onboard_screen/",
            "SplashCubit performs deviceAuth → stores token → navigates to Home or Onboarding screens based on profile completeness.",
        ),
        (
            "Authentication",
            "Email login/register, OTP verification, Google/Apple sign-in, password reset.",
            "lib/presentation/authentication/, lib/data/repositories/auth_repo.dart",
            "AuthScreenCubit calls AuthRepository → saves token to PrefService → navigates to Home.",
        ),
        (
            "Home Shell",
            "Bottom navigation with News, Map, Go Live categories, and Profile tabs.",
            "lib/presentation/main/home/home_screen.dart, home_screen_bloc_cubit.dart",
            "Initializes socket connection, permissions, push token, location updates, and showcase tutorial.",
        ),
        (
            "News Feed",
            "Paginated list of incident/general posts with filters (latest, trending, etc.).",
            "lib/presentation/main/news_screen/",
            "NewsScreenBlocCubit fetches via MainRepository.eventNewsList with infinite scroll pagination.",
        ),
        (
            "News Details",
            "Full post view with video, comments, reactions, reporting, in-area related posts.",
            "lib/presentation/main/news_details_screen/",
            "NewsDetailsScreenBlocCubit loads single post, comments, and handles user interactions.",
        ),
        (
            "Map",
            "Mapbox map displaying nearby event markers within user radius.",
            "lib/presentation/main/map_screen/, lib/core/utils/map_utils.dart",
            "MapScreenCubit fetches area events and renders point annotations on Mapbox.",
        ),
        (
            "Go Live / General Post",
            "Camera-based live video recording, category selection, watermark, compression, and publish.",
            "lib/presentation/main/go_live_screen/, lib/presentation/main/general_go_live/, lib/services/camera_services.dart",
            "GoLiveCubit / GeneralPostCubit upload via MainRepository.postEvent / generalAddPost.",
        ),
        (
            "Send Alert",
            "Send targeted alerts to users within radius.",
            "lib/presentation/main/send_alert/",
            "SendAlertCubit handles alert composition and API submission.",
        ),
        (
            "Search",
            "Search events by hashtag and discover users by username.",
            "lib/presentation/main/search_screen/",
            "SearchScreenBlocCubit queries search APIs with debounced input.",
        ),
        (
            "Chat & Friends",
            "Real-time messaging, friend requests, unread counts via Socket.IO.",
            "lib/presentation/main/chat_screen/, lib/presentation/main/add_friend_screen/, lib/services/socket_service/",
            "SocketService connects with userId query param; events handled via SocketEvents constants.",
        ),
        (
            "Profile & Settings",
            "View/edit profile, notification settings, video quality, blocked users, account deletion.",
            "lib/presentation/main/my_profile_screen/, edit_profile_screen/, settings_screen/, delete_account_screen/",
            "Profile Cubits call MainRepository for CRUD; settings persisted in PrefService.",
        ),
        (
            "Alerts / Notifications",
            "In-app notification list and push notification deep-link handling.",
            "lib/presentation/main/alerts_screen/, lib/services/one_signal_notification_service.dart",
            "OneSignal click listener routes to news details via notification payload.",
        ),
        (
            "Support",
            "Submit and track support requests.",
            "lib/presentation/main/get_support_screen/",
            "GetSupportCubit manages ticket list and creation via support APIs.",
        ),
        (
            "Ads",
            "Remote-config-driven ad placements (splash, news list native, banners).",
            "lib/services/ads_service/, lib/services/remote_config_service/",
            "Firebase Remote Config provides test_ads_json; AdsService loads placements accordingly.",
        ),
    ]
    for name, desc, files, flow in modules:
        add_heading(doc, name, 2)
        add_paragraph(doc, f"Description: {desc}")
        add_paragraph(doc, f"Main Files: {files}")
        add_paragraph(doc, f"Flow: {flow}")

    # 10. UI/UX
    add_heading(doc, "10. UI/UX Structure", 1)
    add_heading(doc, "Navigation", 2)
    add_bullets(
        doc,
        [
            "Navigator 1.0 with named routes (MaterialApp.onGenerateRoute)",
            "Route definitions: lib/routes/app_routes.dart",
            "Navigation helpers: lib/routes/app_navigation.dart (NavigatorRoute class)",
            "Custom page transitions: fade, slide, scale, rotate via PageRouteBuilder",
            "Global navigatorKey: AppFunctions.navigatorKey for context-free navigation",
            "GoRouter / Navigator 2.0: Not used",
        ],
    )
    add_heading(doc, "Theme Management", 2)
    add_bullets(
        doc,
        [
            "Single dark theme defined in lib/core/theme/app_theme.dart (darkTheme)",
            "Color palette: lib/core/theme/colors.dart",
            "Typography: lib/core/theme/text_styles.dart with Roboto, Test Tiempos Headline, and Alice fonts",
            "MaterialApp theme: darkTheme applied globally in main.dart",
        ],
    )
    add_heading(doc, "Responsive Design", 2)
    add_bullets(
        doc,
        [
            "flutter_screenutil with design size 430×990 (portrait phone baseline)",
            "Text scaling locked to 1.0 via MediaQuery textScaler override",
            "Portrait-only orientation enforced in main.dart",
            "Shimmer skeleton loaders for loading states",
        ],
    )

    # 11. Assets
    add_heading(doc, "11. Assets & Resources", 1)
    add_heading(doc, "Images & Icons", 2)
    add_bullets(
        doc,
        [
            "assets/icons/ – SVG and PNG UI icons (search, filter, location, rescue, etc.)",
            "assets/images/ – PNG images (logos, category illustrations, placeholders)",
            "assets/videos/ – bundled video assets if any",
            "assets/animation/ – Lottie JSON (loader.json, success_animation.json)",
            "Generated references: lib/gen/assets.gen.dart (flutter_gen)",
        ],
    )
    add_heading(doc, "Fonts", 2)
    add_bullets(
        doc,
        [
            "Roboto (Light 300 – Black 800) – assets/fonts/roboto/",
            "Test Tiempos Headline – assets/fonts/test_tiempos_headline/",
            "Alice Regular – assets/fonts/alice/",
            "Generated references: lib/gen/fonts.gen.dart",
        ],
    )
    add_heading(doc, "Localization", 2)
    add_paragraph(
        doc,
        "Not Available – No flutter_localizations or ARB files configured. "
        "The intl package is used for date/number formatting only. UI strings are hardcoded in English.",
    )

    # 12. Testing
    add_heading(doc, "12. Testing & QA", 1)
    add_bullets(
        doc,
        [
            "Unit Tests: Not Available – no unit test files beyond default scaffold",
            "Widget Tests: test/widget_test.dart exists but contains outdated counter-app boilerplate (will fail against current app)",
            "Integration Tests: Not Available",
            "Known Test Coverage: ~0% – no meaningful automated test coverage",
            "Manual QA recommended for: auth flows, go-live video upload, map rendering, push notifications, deep links, and payment-free ad flows",
        ],
    )

    # 13. Performance
    add_heading(doc, "13. Performance Considerations", 1)
    add_bullets(
        doc,
        [
            "Deferred initialization: heavy services (ads, crashlytics, camera) load in _deferredInit() after first frame",
            "cached_network_image for efficient image caching",
            "light_compressor_v2 for video compression before upload",
            "Pagination on news feed and notification lists to limit payload size",
            "Firebase Performance Monitoring integrated for production tracing",
            "Release builds: Android R8 minification and resource shrinking enabled",
            "Mapbox Impeller disabled on Android (EnableImpeller=false) for compatibility",
            "Socket reconnection with 5 attempts and 2s delay",
        ],
    )

    # 14. Known Issues
    add_heading(doc, "14. Known Issues / Limitations", 1)
    add_bullets(
        doc,
        [
            "Release Android build uses debug signing config – must be updated for Play Store production",
            "isLiveMode hardcoded to true – environment switching requires code change and rebuild",
            "ApiController Dio instance has a different hardcoded base URL (http://139.59.36.221:8001) than RestClient – potential inconsistency for Dio-only calls",
            "TokenInterceptor in ApiController is commented out",
            "Default widget test is obsolete and does not reflect the application",
            "No localization – English-only UI",
            "Web and desktop platforms not production-ready",
            "Some legacy package IDs in app_constant.dart (com.invoice.maker) appear unused/incorrect",
            "Branch.io TestMode must be manually verified before production release",
            "Guest login flow stores session without full profile – onboarding may be required afterward",
            "Legacy commented reference code from unrelated projects exists in lib/services/deep_linking.dart (BookClubLM/GetX) – should be removed",
        ],
    )

    # 15. Future Improvements
    add_heading(doc, "15. Future Improvements", 1)
    add_bullets(
        doc,
        [
            "Introduce proper environment configuration (flavors or --dart-define) for dev/staging/prod",
            "Add domain layer with use cases to decouple Cubits from repositories",
            "Implement comprehensive unit and widget tests; fix existing widget_test.dart",
            "Migrate to GoRouter or Navigator 2.0 for declarative deep-link routing",
            "Add flutter_localizations and multi-language ARB support",
            "Replace hardcoded secrets with secure CI/CD secret injection",
            "Configure production Android signing and iOS provisioning in CI",
            "Consolidate ApiController Dio usage with RestClient or remove duplicate client",
            "Enable TokenInterceptor for automatic token refresh",
            "Improve error parsing in repositories (avoid brittle string splitting on Failure)",
            "Add integration tests for critical user journeys (login, post alert, receive notification)",
        ],
    )

    # 16. Credentials
    add_heading(doc, "16. Credentials & Access", 1)
    add_paragraph(
        doc,
        "IMPORTANT: Store all credentials securely. Values below are masked for handover documentation. "
        "Request full credentials from the project owner or secrets vault.",
    )
    add_bullets(
        doc,
        [
            "Firebase Project ID: aawaj-52af1",
            "Firebase API Key (Android): AIzaSy********************* (see firebase_options.dart)",
            "OneSignal App ID: 59df38e3-a27a-4f4c-bd65-9e19c39fbff7",
            "Google OAuth Client ID (iOS): 552605622755-****************.apps.googleusercontent.com",
            "Mapbox Access Token: Pass via --dart-define=ACCESS_TOKEN=pk.********",
            "Branch.io: Configured in AndroidManifest.xml and iOS Info.plist (app links: awaazeye.app.link)",
            "Google AdMob (test IDs in Remote Config defaults): ca-app-pub-3940256099942544/...",
            "Bitbucket logic_go_network: Requires authenticated git access to https://bitbucket.org/LogicGoNIkunj/logic_go_network.git",
            "Backend Admin Panel: Not part of this mobile codebase – separate admin web project may exist",
            "Admin Access Notes: Use backend-provided admin credentials for content moderation; not embedded in mobile app",
        ],
    )

    # 17. Handover Notes
    add_heading(doc, "17. Handover Notes", 1)
    add_heading(doc, "Important Developer Notes", 2)
    add_bullets(
        doc,
        [
            "Always run build_runner after pulling changes that affect Freezed models or flutter_gen assets.",
            "Set isLiveMode correctly and verify Branch TestMode=false before any production release.",
            "Verify Firebase Crashlytics is recording errors only when isLiveMode is true.",
            "Mapbox token is mandatory for map functionality – builds without ACCESS_TOKEN will fail on map screens.",
            "Private git dependency logic_go_network must remain accessible for flutter pub get to succeed.",
            "Portrait orientation is enforced – test all screens in portrait only.",
        ],
    )
    add_heading(doc, "Things to Be Careful About", 2)
    add_bullets(
        doc,
        [
            "Do not commit production keystore files or unmasked API secrets to version control.",
            "Changing baseUrl or isLiveMode affects both REST and Socket connections simultaneously.",
            "OneSignal and Branch deep-link handlers depend on app initialization order – notification clicks before init are queued in notificationPayload.",
            "Video upload flows depend on camera permissions, compression, and network – test on real devices.",
            "401 token expiry triggers global logout – ensure refresh token flow is implemented before extending session duration.",
            "Remote Config ad JSON structure must match AdsJsonModel parsing – invalid JSON will break ad loading.",
        ],
    )

    # Footer
    doc.add_page_break()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.add_run("— End of Document —\n").bold = True
    footer_p.add_run(f"Awaaz Flutter Mobile App Project Handover | {TODAY}")

    return doc


def main():
    doc = build_document()
    doc.save(DOCX_PATH)
    print(f"DOCX created: {DOCX_PATH}")

    try:
        from docx2pdf import convert
        convert(str(DOCX_PATH), str(PDF_PATH))
        print(f"PDF created: {PDF_PATH}")
    except Exception as e:
        print(f"docx2pdf conversion failed: {e}")
        print("Attempting ReportLab PDF fallback...")
        create_pdf_fallback()


def create_pdf_fallback():
    """Fallback PDF using reportlab if Word is unavailable."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Heading1"],
        fontSize=18,
        textColor=colors.HexColor("#006666"),
        spaceAfter=12,
    )
    h2_style = ParagraphStyle(
        "CustomH2",
        parent=styles["Heading2"],
        fontSize=14,
        textColor=colors.HexColor("#006666"),
        spaceAfter=8,
    )

    story = []
    story.append(Paragraph("Awaaz Flutter Mobile App – Project Handover Document", title_style))
    story.append(Paragraph(f"Prepared: {TODAY} | Version 1.0.13+14", styles["Normal"]))
    story.append(Spacer(1, 0.3 * inch))
    story.append(
        Paragraph(
            "This PDF is a condensed export. For the complete handover document with all 17 sections, "
            "refer to the companion DOCX file: Awaaz_Flutter_Mobile_App_Project_Handover_Document.docx",
            styles["Normal"],
        )
    )
    story.append(PageBreak())

    sections = [
        ("1. Project Overview", "Awaaz is a real-time city alert mobile app (package: eagle_eye) for Android and iOS."),
        ("2. Technology Stack", "Flutter 3.44.1, Dart 3.12.1, BLoC/Cubit + Freezed, Dio, Mapbox, Firebase, OneSignal."),
        ("3. Architecture", "Layered Presentation–Data–Core pattern with feature-based organization."),
        ("4. Folder Structure", "lib/core, lib/presentation, lib/data, lib/services, lib/routes. No domain/ or features/ layer."),
        ("5. Setup", "flutter pub get → build_runner → flutter run --dart-define=ACCESS_TOKEN=..."),
        ("6. Build", "flutter build apk/appbundle/ios --release --dart-define=ACCESS_TOKEN=..."),
        ("7. State Management", "Global MultiBlocProvider with per-feature Cubits and Freezed states."),
        ("8. API", "REST via logic_go_network; production https://awaazeye.com; Bearer token auth."),
        ("9. Modules", "Auth, News, Map, Go Live, Chat, Profile, Alerts, Support, Ads."),
        ("10. UI/UX", "Navigator 1.0, dark theme, ScreenUtil responsive scaling."),
        ("11. Assets", "assets/icons, images, fonts, animation. No localization."),
        ("12. Testing", "Minimal – outdated widget test only. Coverage ~0%."),
        ("13. Performance", "Deferred init, caching, compression, pagination, R8 shrinking."),
        ("14. Known Issues", "Debug signing on release, hardcoded env toggle, no tests, English only."),
        ("15. Future", "Flavors, tests, localization, GoRouter, production signing."),
        ("16. Credentials", "Firebase aawaj-52af1, OneSignal, Mapbox, Branch – see DOCX for masked details."),
        ("17. Handover Notes", "Verify Branch TestMode, Crashlytics, Mapbox token, logic_go_network access."),
    ]
    for title, body in sections:
        story.append(Paragraph(title, h2_style))
        story.append(Paragraph(body, styles["Normal"]))
        story.append(Spacer(1, 0.15 * inch))

    doc.build(story)
    print(f"PDF fallback created: {PDF_PATH}")


if __name__ == "__main__":
    main()
